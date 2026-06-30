import os
import pandas as pd
import docx

def create_mocks():
    mock_dir = "mock_inputs"
    resumes_dir = os.path.join(mock_dir, "resumes")
    os.makedirs(resumes_dir, exist_ok=True)

                      
    recruiter_data = {
        "name": ["Akash Chandra", "Priya Sharma", "Rahul Verma"],
        "email": ["akash.chandra@gmail.com", "priya.sharma@yahoo.com", "rahul.verma@gmail.com"],
        "phone": ["9876543210", "9876543211", "9876543212"],
        "city": ["Bengaluru", "Pune", "Gurugram"],
        "country": ["IN", "IN", "IN"]
    }
    pd.DataFrame(recruiter_data).to_csv(os.path.join(mock_dir, "recruiter.csv"), index=False)

                   
    github_data = {
        "name": ["Akash Chandra", "Priya Sharma", "Rahul Verma"],
        "email": ["akash.chandra@gmail.com", "priya.sharma@yahoo.com", "rahul.verma@gmail.com"],
        "github_link": [
            "https://github.com/akash-chandra-105",
            "https://github.com/priya-sharma-dev",
            "https://github.com/rverma-backend"
        ]
    }
    pd.DataFrame(github_data).to_csv(os.path.join(mock_dir, "github.csv"), index=False)

                                 
    doc1 = docx.Document()
    doc1.add_heading("Akash Chandra", 0)
    doc1.add_paragraph("Email: akash.chandra@gmail.com, akash.c@infosys.com")
    doc1.add_paragraph("Phone: +919876543210")
    doc1.add_paragraph("LinkedIn: linkedin.com/in/akash-chandra")
    doc1.add_paragraph("Skills: Python, JavaScript, React, PostgreSQL, Node.js, Docker, AWS")
    doc1.add_paragraph("Experience: Senior Software Engineer at Infosys (2022-03 to Present)")
    doc1.add_paragraph("Experience: Software Engineer at TCS (2019-06 to 2022-02)")
    doc1.add_paragraph("Education: VIT University, B.Tech in Computer Science, 2017")
    doc1.save(os.path.join(resumes_dir, "akash_resume.docx"))

                                
    doc2 = docx.Document()
    doc2.add_heading("Priya Sharma", 0)
    doc2.add_paragraph("Email: priya.sharma@yahoo.com, priya.s@wipro.com")
    doc2.add_paragraph("Phone: +919876543211")
    doc2.add_paragraph("LinkedIn: linkedin.com/in/priya-sharma")
    doc2.add_paragraph("Skills: JavaScript, TypeScript, React, Next.js, Redux, HTML, CSS")
    doc2.add_paragraph("Experience: Frontend Architect at Wipro (2023-01 to Present)")
    doc2.add_paragraph("Education: Pune University, B.E. in IT, 2020")
    doc2.save(os.path.join(resumes_dir, "priya_resume.docx"))

    print(f"Mock inputs generated successfully in: {os.path.abspath(mock_dir)}")

if __name__ == "__main__":
    create_mocks()
